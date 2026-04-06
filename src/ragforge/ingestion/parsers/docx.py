from __future__ import annotations

from pathlib import Path

from docx import Document as DocxDocument

from ragforge.core.schemas import Document


def load_docx_file(file_path: str | Path) -> Document | None:
    """Load a `.docx` file into a `Document` using python-docx."""
    path = Path(file_path)
    docx = DocxDocument(path)
    paragraphs = [paragraph.text.strip() for paragraph in docx.paragraphs]
    text = "\n\n".join(paragraph for paragraph in paragraphs if paragraph)
    if not text.strip():
        return None

    return Document(
        doc_id=path.stem,
        text=text.strip(),
        metadata={
            "source": str(path),
            "filename": path.name,
            "file_type": "docx",
            "parser": "python-docx",
        },
    )
